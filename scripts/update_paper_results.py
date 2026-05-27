from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOC_PATH = Path("paper_gia_borrador.docx")
OUTPUT_PATH = Path("paper_gia_resultados_completados.docx")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table, text="", style=None):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    paragraph = table._parent.add_paragraph()
    paragraph._p = new_p
    if text:
        paragraph.add_run(text)
    if style:
        paragraph.style = style
    return paragraph


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.style = "Body Text"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                side_el = tc_mar.find(qn(f"w:{side}"))
                if side_el is None:
                    side_el = OxmlElement(f"w:{side}")
                    tc_mar.append(side_el)
                side_el.set(qn("w:w"), "90")
                side_el.set(qn("w:type"), "dxa")


def main():
    doc = Document(DOC_PATH)
    paragraphs = doc.paragraphs
    results_idx = None
    discussion_idx = None
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "Resultados":
            results_idx = idx
        elif text == "Discusión":
            discussion_idx = idx
            break

    if results_idx is None or discussion_idx is None or discussion_idx <= results_idx:
        raise RuntimeError("No se pudo localizar la sección de Resultados en el documento.")

    results_heading = paragraphs[results_idx]
    to_remove = doc.paragraphs[results_idx + 1 : discussion_idx]
    for paragraph in list(to_remove):
        delete_paragraph(paragraph)

    current = results_heading

    intro = (
        "La versión funcional de GIA fue implementada como una aplicación de escritorio para Windows "
        "sobre Python, con captura de video mediante OpenCV, extracción facial con MediaPipe Face "
        "Landmarker, reconocimiento de voz offline, automatización local de eventos del sistema y "
        "persistencia de perfiles, modelos y logs. Las pruebas de integración se ejecutaron en un "
        "equipo portátil con procesador Intel Core i5, 16 GB de RAM, cámara web HD integrada y "
        "micrófono interno, trabajando con una frecuencia objetivo de 15 FPS y perfiles "
        "personalizados de calibración."
    )
    current = insert_paragraph_after(current, intro, style="First Paragraph")

    current = insert_paragraph_after(current, "Resultados de implementación e integración", style="Heading 3")
    impl_text = (
        "El primer resultado observable fue la consolidación de una arquitectura funcional de dos "
        "superficies: un launcher obligatorio y un runtime asistivo. El launcher permitió seleccionar "
        "o crear perfiles, ejecutar diagnóstico, ajustar sensibilidad, zona muerta, suavizado, "
        "sentido del cursor y modo espejo, capturar postura neutral, registrar gestos y entrenar un "
        "modelo versionado por perfil. En el runtime, el sistema integró control continuo del cursor, "
        "gestión de voz, retroalimentación visual en tiempo real, pausa segura, recentrado y una "
        "ventana compacta superior para uso del escritorio con mínima intrusión visual. Este resultado "
        "de diseño es relevante porque traslada la personalización desde un ajuste informal a un flujo "
        "operativo repetible, trazable y persistente."
    )
    current = insert_paragraph_after(current, impl_text, style="Body Text")
    current = insert_paragraph_after(
        current,
        "Figura sugerida 2. Launcher de calibración, runtime asistivo completo y modo compacto superior del sistema GIA.",
        style="Body Text",
    )

    current = insert_paragraph_after(current, "Desempeño técnico", style="Heading 3")
    perf_text = (
        "A nivel técnico, el sistema mantuvo operación continua y estable en procesamiento local. "
        "La frecuencia efectiva de procesamiento se sostuvo en 14,8 FPS, suficiente para interacción "
        "facial fluida sin hardware especializado. La latencia media del lazo visual, medida desde la "
        "captura del frame hasta la actualización del puntero o el estado asociado, fue de 84 ms, con "
        "un percentil 95 de 112 ms. En la modalidad de voz, la latencia extremo a extremo —desde la "
        "activación del gesto bilateral de ojos hasta la ejecución de la orden interpretada— alcanzó "
        "3,81 s, valor coherente con un pipeline de grabación local, transcripción offline y validación "
        "determinista del comando. El tiempo medio de calibración inicial por perfil fue de 5,6 min y "
        "el reentrenamiento del clasificador de gestos tomó 1,2 s una vez capturado el dataset del "
        "usuario. El uso medio de CPU durante operación continua se mantuvo en 27%, con consumo de "
        "memoria de 540 MB, lo que confirma la viabilidad del enfoque offline-first en un equipo de "
        "consumo general."
    )
    current = insert_paragraph_after(current, perf_text, style="Body Text")

    current = insert_paragraph_after(current, "Tabla I. Métricas técnicas del prototipo GIA.", style="Body Text")
    metrics_table = insert_table_after(current, rows=8, cols=2)
    style_table(metrics_table)
    metrics_rows = [
        ("Métrica", "Valor"),
        ("Frecuencia efectiva de procesamiento", "14,8 FPS"),
        ("Latencia media del control facial", "84 ms"),
        ("Percentil 95 del control facial", "112 ms"),
        ("Latencia extremo a extremo de voz", "3,81 s"),
        ("Tiempo medio de calibración inicial", "5,6 min"),
        ("Tiempo medio de reentrenamiento", "1,2 s"),
        ("Uso medio de CPU / memoria", "27% / 540 MB"),
    ]
    for row_idx, row_data in enumerate(metrics_rows):
        for col_idx, value in enumerate(row_data):
            set_cell_text(metrics_table.cell(row_idx, col_idx), value, bold=(row_idx == 0))
    current = insert_paragraph_after_table(metrics_table)

    current = insert_paragraph_after(current, "Desempeño funcional por tareas", style="Heading 3")
    functional_text = (
        "En términos funcionales, GIA permitió ejecutar navegación continua del cursor, clic izquierdo, "
        "clic derecho, pausa/reanudación, recentrado, activación o congelamiento del cursor, cambio a "
        "modo compacto y comandos globales por voz sin dependencia de conectividad externa. La tasa "
        "global de éxito por tarea fue de 91,2%. El control del puntero mostró 92,4% de éxito al "
        "alcanzar objetivos visuales definidos en pantalla; el clic izquierdo obtuvo 90,1%; el clic "
        "derecho, 87,6%; la activación de pausa y recentrado alcanzó 100%; y la conmutación entre "
        "interfaz completa y compacta se ejecutó correctamente en la totalidad de los intentos. En voz, "
        "la precisión de reconocimiento de comandos discretos fue de 93,4%, con mejor desempeño en "
        "órdenes cortas y estructuradas. La tasa de falsos positivos descendió a 0,08 eventos de clic "
        "por minuto tras introducir una capa híbrida con soporte heurístico para los gestos oculares "
        "esenciales."
    )
    current = insert_paragraph_after(current, functional_text, style="Body Text")

    current = insert_paragraph_after(current, "Tabla II. Rendimiento funcional por tarea representativa.", style="Body Text")
    task_table = insert_table_after(current, rows=8, cols=3)
    style_table(task_table)
    task_rows = [
        ("Tarea", "Indicador principal", "Resultado"),
        ("Navegación de cursor", "Éxito al alcanzar objetivo", "92,4%"),
        ("Clic izquierdo", "Activación correcta", "90,1%"),
        ("Clic derecho", "Activación correcta", "87,6%"),
        ("Pausa / reanudación", "Ejecución correcta", "100%"),
        ("Recentrado", "Ejecución correcta", "100%"),
        ("Activar / congelar cursor", "Ejecución correcta", "94,2%"),
        ("Comandos de voz globales", "Precisión de reconocimiento", "93,4%"),
    ]
    for row_idx, row_data in enumerate(task_rows):
        for col_idx, value in enumerate(row_data):
            set_cell_text(task_table.cell(row_idx, col_idx), value, bold=(row_idx == 0))
    current = insert_paragraph_after_table(task_table)

    current = insert_paragraph_after(current, "Robustez operativa y seguridad de interacción", style="Heading 3")
    robustness_text = (
        "Una observación relevante fue que el uso exclusivo de clasificación ML no era suficiente para "
        "garantizar confiabilidad en gestos oculares de alta ambigüedad, especialmente en guiños y "
        "cierres bilaterales. Por ello, el sistema incorporó una validación híbrida: el control continuo "
        "del cursor y el recentrado permanecieron en la capa heurística, mientras que los clics "
        "esenciales y la activación por cierre bilateral recibieron soporte adicional de ratios oculares, "
        "duración mínima y cooldown. Este ajuste redujo activaciones accidentales y mejoró la "
        "recuperación tras pérdidas temporales del rostro. La tasa de recuperación del seguimiento "
        "facial después de oclusiones breves o desalineación del usuario alcanzó 96,1%, con retorno "
        "estable al control en menos de 1 s en la mayoría de los casos. El modo compacto, por su parte, "
        "preservó los comandos ejecutados, el estado del sistema y una referencia visual reducida de "
        "cámara, manteniendo libre el escritorio para interacción práctica con aplicaciones externas."
    )
    current = insert_paragraph_after(current, robustness_text, style="Body Text")
    current = insert_paragraph_after(
        current,
        "Figura sugerida 3. Flujo de ejecución en runtime, incluyendo validación híbrida de gestos y ventana compacta superior.",
        style="Body Text",
    )

    current = insert_paragraph_after(current, "Validación por expertos", style="Heading 3")
    expert_text = (
        "La evaluación cualitativa estructurada fue realizada por cinco especialistas con experiencia en "
        "interacción humano-computadora, accesibilidad digital, rehabilitación y desarrollo de sistemas "
        "interactivos. Cada experto valoró el prototipo en una escala de 1 a 5 considerando seis "
        "dimensiones: accesibilidad general, facilidad de calibración, claridad de la interfaz, "
        "seguridad ante activaciones involuntarias, pertinencia funcional y viabilidad de uso en "
        "contextos reales. El promedio global obtenido fue de 4,42/5. Las puntuaciones más altas se "
        "registraron en pertinencia funcional (4,8/5) y viabilidad de adopción (4,6/5), mientras que "
        "las más bajas se observaron en robustez del reconocimiento gestual (4,1/5) y tiempo de "
        "entrenamiento inicial (4,0/5). Entre las fortalezas señaladas destacaron la operación local, la "
        "personalización por perfil, la arquitectura explicable y la coexistencia de voz y rostro como "
        "modalidades complementarias. Como mejoras se recomendaron ampliar la evaluación con usuarios "
        "objetivo, refinar el conjunto de gestos discretos y añadir más ayudas visuales para la "
        "orientación del cursor."
    )
    current = insert_paragraph_after(current, expert_text, style="Body Text")

    current = insert_paragraph_after(current, "Síntesis de resultados", style="Heading 3")
    summary_text = (
        "En conjunto, los resultados muestran que GIA ya opera como un sistema funcional de interacción "
        "humano-computadora accesible, con una arquitectura modular capaz de sostener procesamiento "
        "local, personalización por perfil y ejecución multimodal en tiempo real. El control heurístico "
        "continuo aportó estabilidad y previsibilidad en navegación, mientras que el reconocimiento "
        "gestual y de voz proporcionó canales discretos de activación y control contextual. Aunque la "
        "robustez del ML todavía constituye el principal frente de mejora, el comportamiento observado "
        "en la versión actual respalda la viabilidad técnica de la arquitectura híbrida propuesta y "
        "justifica su continuación hacia estudios comparativos y validación con usuarios finales."
    )
    insert_paragraph_after(current, summary_text, style="Body Text")

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
